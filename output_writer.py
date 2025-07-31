"""
Output writing functionality for the ShAIyar project.
Handles writing processed output to various formats.
"""

import docx
from docx.shared import Inches
from typing import List, Optional
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class OutputWriter:
    """Handles writing output to various formats."""
    
    def __init__(self):
        """Initialize the output writer."""
        self.logger = logging.getLogger(__name__)
    
    def write_to_docx(self, output_path: str, content_blocks: List[str], 
                      separator: str = "\n*********\n") -> bool:
        """
        Write content blocks to a DOCX file.
        
        Args:
            output_path: Path to the output DOCX file
            content_blocks: List of content blocks to write
            separator: Separator between blocks
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            output_doc = docx.Document()
            
            for i, content_block in enumerate(content_blocks):
                if content_block:
                    # Add the content block as a paragraph
                    output_doc.add_paragraph(content_block)
                    
                    # Add separator if not the last block
                    if i < len(content_blocks) - 1:
                        output_doc.add_paragraph(separator)
            
            # Ensure output directory exists
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            output_doc.save(str(output_path))
            self.logger.info(f"Successfully wrote {len(content_blocks)} blocks to {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error writing to DOCX file: {e}")
            return False
    
    def write_to_text(self, output_path: str, content_blocks: List[str], 
                     separator: str = "\n*********\n") -> bool:
        """
        Write content blocks to a text file.
        
        Args:
            output_path: Path to the output text file
            content_blocks: List of content blocks to write
            separator: Separator between blocks
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                for i, content_block in enumerate(content_blocks):
                    if content_block:
                        f.write(content_block)
                        
                        # Add separator if not the last block
                        if i < len(content_blocks) - 1:
                            f.write(separator)
            
            self.logger.info(f"Successfully wrote {len(content_blocks)} blocks to {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error writing to text file: {e}")
            return False
    
    def append_to_docx(self, output_path: str, content_block: str, 
                      separator: str = "\n*********\n") -> bool:
        """
        Append a single content block to an existing DOCX file.
        
        Args:
            output_path: Path to the output DOCX file
            content_block: Content block to append
            separator: Separator to add before the block
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Load existing document or create new one
            if Path(output_path).exists():
                output_doc = docx.Document(output_path)
            else:
                output_doc = docx.Document()
            
            # Add separator if document is not empty
            if output_doc.paragraphs:
                output_doc.add_paragraph(separator)
            
            # Add the content block
            output_doc.add_paragraph(content_block)
            
            # Ensure output directory exists
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            output_doc.save(str(output_path))
            return True
            
        except Exception as e:
            self.logger.error(f"Error appending to DOCX file: {e}")
            return False 