"""
DOCX file reading and text extraction functionality.
Handles reading and extracting text blocks from .docx files.
"""

import docx
from typing import Generator, Optional
import logging
import os

logger = logging.getLogger(__name__)


class DocxReader:
    """Handles reading and extracting text from .docx files."""
    
    def __init__(self):
        """Initialize the DOCX reader."""
        self.logger = logging.getLogger(__name__)
    
    def read_docx_text_blocks(self, file_path: str) -> Generator[Optional[str], None, None]:
        """
        Read a DOCX file and yield text blocks separated by empty lines.
        
        Args:
            file_path: Path to the input DOCX file
            
        Yields:
            Optional[str]: Text blocks or None if error occurs
        """
        try:
            doc = docx.Document(file_path)
            text_block = ""
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                
                if text.strip() == "":  # Check for empty paragraph
                    if text_block:  # Yield the block if it's not empty
                        yield text_block.strip()
                        text_block = ""  # Reset the block
                else:
                    text_block += text + "\n"  # Add text to the block with a newline
            
            # Yield the last block if it exists after the loop
            if text_block:
                yield text_block.strip()
                
        except FileNotFoundError:
            self.logger.error(f"File not found at {file_path}")
            yield None
        except Exception as e:
            self.logger.error(f"Error reading DOCX file: {e}")
            yield None
    
    def get_document_info(self, file_path: str) -> dict:
        """
        Get basic information about the DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            dict: Document information including paragraph count, etc.
        """
        try:
            doc = docx.Document(file_path)
            return {
                "paragraph_count": len(doc.paragraphs),
                "file_path": file_path,
                "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }
        except Exception as e:
            self.logger.error(f"Error getting document info: {e}")
            return {} 