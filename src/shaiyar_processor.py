"""
Main processing module for the ShAIyar project.
Orchestrates the entire text processing pipeline.
"""

import time
import logging
from typing import List, Optional
from pathlib import Path
from docx import Document

from config import Config
from docx_reader import DocxReader
from text_processor import TextProcessor
from llm_interface import LLMFactory
from output_writer import OutputWriter


class ShAIyarProcessor:
    """Main processor class for the ShAIyar project."""
    
    def __init__(self, config: Config):
        """
        Initialize the ShAIyar processor.
        
        Args:
            config: Configuration object
        """
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Initialize components
        self.docx_reader = DocxReader()
        self.text_processor = TextProcessor()
        self.output_writer = OutputWriter()
        
        # Initialize LLM interface
        llm_config = {
            'model_name': config.llm_config.model_name,
            'api_key': config.llm_config.api_key,
            'temperature': config.llm_config.temperature,
            'max_tokens': config.llm_config.max_tokens,
            'max_retries': config.processing_config.max_retries
        }
        
        self.llm_interface = LLMFactory.create_llm_interface(
            config.llm_config.provider, 
            llm_config
        )
        
        # Load system message
        self.system_message = self._load_system_message()
    
    def _load_system_message(self) -> str:
        """
        Load the system message from file.
        
        Returns:
            str: System message content
        """
        try:
            with open(self.config.file_config.system_message_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            self.logger.error(f"Error loading system message: {e}")
            return ""
    
    def _process_text_block(self, text_block: str, block_index: int) -> Optional[str]:
        """
        Process a single text block through the LLM.
        
        Args:
            text_block: Text block to process
            block_index: Index of the block for logging
            
        Returns:
            Optional[str]: Processed text block or None if error
        """
        try:
            # Clean the text block
            cleaned_text = self.text_processor.clean_text(text_block)
            
            if not cleaned_text:
                self.logger.warning(f"Block {block_index} is empty after cleaning")
                return None
            
            # Analyze the text block
            analysis = self.text_processor.analyze_text_block(cleaned_text)
            self.logger.info(f"Processing block {block_index}: {analysis}")
            
            # Create messages for LLM
            messages = self.llm_interface.create_messages(
                self.system_message, 
                cleaned_text
            )
            
            # Invoke LLM
            result = self.llm_interface.invoke(messages)
            
            self.logger.info(f"Successfully processed block {block_index}")
            return result
            
        except Exception as e:
            self.logger.error(f"Error processing block {block_index}: {e}")
            return None
    
    def process(self) -> bool:
        """
        Process the entire document, showing progress and saving incrementally.
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            # Validate configuration
            self.config.validate()
            
            # Count total blocks first
            text_blocks = list(self.docx_reader.read_docx_text_blocks(
                self.config.file_config.input_docx_path
            ))
            total_blocks = len([b for b in text_blocks if b is not None])
            
            self.logger.info(f"Starting processing of {total_blocks} blocks")
            
            # Create an empty document to write to.
            output_doc = Document()
            is_first_block_written = True
            
            # Process blocks with progress
            successful_blocks = 0
            
            for i, text_block in enumerate(text_blocks):
                if text_block is None:
                    continue
                
                self.logger.info(f"Processing block {i+1}/{total_blocks}")
                
                processed_block = self._process_text_block(text_block, i)
                if processed_block:
                    # Add separator if this is not the first block
                    if not is_first_block_written and self.config.file_config.separator:
                        output_doc.add_paragraph(self.config.file_config.separator)
                    
                    # Add processed block to the document
                    output_doc.add_paragraph(processed_block)
                    is_first_block_written = False
                    
                    # Save the document after each block
                    output_doc.save(self.config.file_config.output_docx_path)
                    self.logger.info(f"Saved progress for block {i+1} to output file.")
                    
                    successful_blocks += 1
                
                # Add delay between requests
                if i < len(text_blocks) - 1:
                    time.sleep(self.config.processing_config.delay_between_requests)
            
            self.logger.info(f"Successfully processed {successful_blocks}/{total_blocks} blocks")
            return True
                
        except Exception as e:
            self.logger.error(f"Error processing document: {e}")
            return False 